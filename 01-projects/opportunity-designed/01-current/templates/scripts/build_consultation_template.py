from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "Opportunity_Designed_Consultation_Script_Template.docx"

NAVY = "17324D"
BLUE = "315B78"
GOLD = "C59A55"
INK = "25313B"
MUTED = "687681"
PALE_BLUE = "F2F6F8"
PALE_GOLD = "FBF7EF"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=160, start=220, bottom=160, end=220):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    col = OxmlElement("w:gridCol")
    col.set(qn("w:w"), str(width_dxa))
    grid.append(col)
    for row in table.rows:
        cell = row.cells[0]
        tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            cell._tc.get_or_add_tcPr().append(tc_w)
        tc_w.set(qn("w:w"), str(width_dxa))
        tc_w.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def set_run(run, size=11, bold=False, color=INK, italic=False, font="Aptos"):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, before=0, after=0, line=1.15, keep=False):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.keep_together = keep


def add_copy_block(doc, label, guidance, text, fill):
    heading = doc.add_paragraph()
    style_paragraph(heading, before=12, after=2, line=1.0, keep=True)
    r = heading.add_run(label.upper())
    set_run(r, size=10, bold=True, color=BLUE)

    note = doc.add_paragraph()
    style_paragraph(note, after=6, line=1.0, keep=True)
    r = note.add_run(guidance)
    set_run(r, size=9.5, italic=True, color=MUTED)

    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    remove_table_borders(table)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=210, start=260, bottom=210, end=260)
    p = cell.paragraphs[0]
    style_paragraph(p, after=0, line=1.22, keep=True)
    r = p.add_run(text)
    set_run(r, size=11, color=INK)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
style_paragraph(hp, after=0, line=1.0)
run = hp.add_run("OPPORTUNITY DESIGNED")
set_run(run, size=9, bold=True, color=BLUE)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
style_paragraph(fp, after=0, line=1.0)
run = fp.add_run("Client communication template  •  Editable reference")
set_run(run, size=8.5, color=MUTED)

kicker = doc.add_paragraph()
style_paragraph(kicker, after=5, line=1.0)
r = kicker.add_run("CLIENT COMMUNICATION TEMPLATE")
set_run(r, size=9.5, bold=True, color=GOLD)

title = doc.add_paragraph()
style_paragraph(title, after=4, line=1.0, keep=True)
r = title.add_run("Diagnostic Fee Credit Language")
set_run(r, size=24, bold=True, color=NAVY)

subtitle = doc.add_paragraph()
style_paragraph(subtitle, after=12, line=1.1, keep=True)
r = subtitle.add_run("Copy-ready wording for consultations and follow-up emails")
set_run(r, size=11.5, color=MUTED)

rule = doc.add_paragraph()
style_paragraph(rule, after=9, line=1.0)
p_pr = rule._p.get_or_add_pPr()
p_bdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "10")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), GOLD)
p_bdr.append(bottom)
p_pr.append(p_bdr)

intro = doc.add_paragraph()
style_paragraph(intro, after=4, line=1.15)
r = intro.add_run("Purpose: ")
set_run(r, size=10.5, bold=True, color=NAVY)
r = intro.add_run("Explain the possible invoice credit clearly while preserving the project deposit terms and written-approval requirement.")
set_run(r, size=10.5, color=INK)

consultation_text = (
    "“If this diagnostic leads us into a larger engagement, I sometimes offer the $450 as an invoice credit "
    "on qualifying projects of $5,000 or more. The project deposit would still be due in full. The credit would "
    "be applied to a later invoice after the deposit is paid. It isn’t a refund or cash-back payment, and any "
    "credit would be confirmed in the written proposal.”"
)

email_text = (
    "One additional note: if the diagnostic leads to a larger Opportunity Designed engagement of $5,000 or more, "
    "the $450 diagnostic fee may be eligible for an invoice credit. The project deposit would remain due in full, "
    "with any approved credit applied to a later project invoice. This is not a refund or cash-back offer. Any "
    "credit and its application would be confirmed in the project proposal."
)

add_copy_block(
    doc,
    "Consultation script",
    "Use verbally when explaining the diagnostic-to-project pathway.",
    consultation_text,
    PALE_BLUE,
)

add_copy_block(
    doc,
    "Email version",
    "Paste into a follow-up email, proposal-prep message, or diagnostic recap.",
    email_text,
    PALE_GOLD,
)

note = doc.add_paragraph()
style_paragraph(note, before=12, after=0, line=1.15, keep=True)
r = note.add_run("Internal reminder: ")
set_run(r, size=9.5, bold=True, color=NAVY)
r = note.add_run("Confirm eligibility and the invoice where the credit will be applied in the written project proposal.")
set_run(r, size=9.5, color=MUTED)

doc.core_properties.title = "Opportunity Designed - Diagnostic Fee Credit Language"
doc.core_properties.subject = "Consultation and email script template"
doc.core_properties.author = "Opportunity Designed"
doc.save(OUTPUT)
print(OUTPUT)
